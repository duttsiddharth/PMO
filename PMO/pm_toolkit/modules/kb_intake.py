"""Knowledge Intake (KB): populate project data from existing documents.

Upload an RFP, SOW, solution design, charter or similar document (PDF, DOCX,
TXT, MD). The toolkit extracts the text, proposes values for the project
charter fields (business case, scope, objectives, deliverables, assumptions,
constraints, success criteria) plus candidate milestones and risks, and lets
you review/edit everything before applying it to a new or existing project.

Extraction uses the Anthropic API when ANTHROPIC_API_KEY is configured;
otherwise a deterministic heading-based parser keeps the feature working
fully offline (same design principle as the rest of core/ai.py).
"""
from __future__ import annotations

import datetime as _dt

import streamlit as st

import config
from core import models as m
from core.ai import extract_project_fields
from modules.common import get_session, list_projects, section_title

CHARTER_FIELDS = [
    ("business_case", "Business Case"),
    ("scope", "Scope"),
    ("objectives", "Objectives"),
    ("deliverables", "Deliverables"),
    ("assumptions", "Assumptions"),
    ("constraints", "Constraints"),
    ("success_criteria", "Success Criteria"),
]


# --------------------------------------------------------------------------
# Text extraction per file type
# --------------------------------------------------------------------------
def _read_pdf(file) -> str:
    from pypdf import PdfReader
    reader = PdfReader(file)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(file) -> str:
    import docx
    doc = docx.Document(file)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_text(upload) -> str:
    name = upload.name.lower()
    try:
        if name.endswith(".pdf"):
            return _read_pdf(upload)
        if name.endswith(".docx"):
            return _read_docx(upload)
        return upload.read().decode("utf-8", errors="replace")
    except Exception as exc:
        st.error(f"Could not read {upload.name}: {exc}")
        return ""


# --------------------------------------------------------------------------
# Apply extracted data
# --------------------------------------------------------------------------
def _apply(fields: dict, milestones: list[str], risks: list[str],
           target: str, existing_id: int | None, new_code: str,
           overwrite: bool):
    s = get_session()
    if target == "new":
        p = m.Project(code=new_code or f"PRJ-KB-{int(_dt.datetime.now().timestamp())%100000}",
                      name=fields.get("name") or "Imported Project",
                      customer=fields.get("customer") or "",
                      project_type=fields.get("project_type") or "",
                      region=fields.get("region") or "",
                      status="Presales", health="Green")
        s.add(p)
        s.flush()
    else:
        p = s.get(m.Project, existing_id)
        if p is None:
            st.error("Target project not found.")
            return None

    for key, _label in CHARTER_FIELDS:
        val = (fields.get(key) or "").strip()
        if val and (overwrite or not (getattr(p, key) or "").strip()):
            setattr(p, key, val)

    for ms in milestones:
        ms = ms.strip()
        if ms and not any(x.name == ms for x in p.milestones):
            p.milestones.append(m.Milestone(name=ms[:200], status="Pending"))

    for rk in risks:
        rk = rk.strip()
        if rk and not any(x.title == rk for x in p.raid):
            p.raid.append(m.RaidItem(category="Risk", title=rk[:200],
                                     owner=p.pm_name or "PM", severity="Medium",
                                     probability=3, impact=3, status="Open",
                                     mitigation="Imported from KB document — assess."))
    s.commit()
    return p


# --------------------------------------------------------------------------
# Page
# --------------------------------------------------------------------------
def render():
    section_title("Knowledge Intake (KB)",
                  "Populate project fields from an RFP, SOW, solution design or charter")

    if st.session_state.get("role", "Customer") not in ("Admin", "PM"):
        st.info("Knowledge intake is available to Admin and PM roles.")
        return

    engine = "AI-assisted (Anthropic)" if config.ANTHROPIC_API_KEY else \
             "Offline heading parser (set ANTHROPIC_API_KEY for AI extraction)"
    st.caption(f"Extraction engine: **{engine}**")

    c1, c2 = st.columns([2, 1])
    upload = c1.file_uploader("Document", type=["pdf", "docx", "txt", "md"],
                              help="RFP, SOW, solution design, charter, proposal …")
    doc_type = c2.selectbox("Document type",
                            ["RFP", "SOW", "Solution Design", "Charter", "Proposal", "Other"])

    if upload and st.button("Extract fields from document", type="primary"):
        with st.spinner("Reading and analysing the document…"):
            text = _extract_text(upload)
            if not text.strip():
                st.error("No readable text found (scanned/image-only PDFs need OCR first).")
            else:
                st.session_state["kb_result"] = extract_project_fields(text, doc_type)
                st.session_state["kb_source"] = f"{upload.name} ({doc_type})"

    result = st.session_state.get("kb_result")
    if not result:
        st.info("Upload a document and click **Extract fields** to begin. Nothing is "
                "written to the database until you review and apply.")
        return

    st.success(f"Extracted from {st.session_state.get('kb_source', 'document')} "
               f"— engine: {result.get('_engine', '?')}. Review and edit below, "
               "then apply.")

    # --- Review & edit -----------------------------------------------------
    with st.form("kb_review"):
        i1, i2, i3, i4 = st.columns(4)
        name = i1.text_input("Project name", result.get("name", ""))
        customer = i2.text_input("Customer", result.get("customer", ""))
        ptype = i3.selectbox("Project type", config.PROJECT_TYPES,
                             index=config.PROJECT_TYPES.index(result["project_type"])
                             if result.get("project_type") in config.PROJECT_TYPES else 0)
        region = i4.text_input("Region", result.get("region", ""))

        edited = {}
        cols = st.columns(2)
        for n, (key, label) in enumerate(CHARTER_FIELDS):
            edited[key] = cols[n % 2].text_area(label, result.get(key, ""), height=110)

        ms_text = st.text_area("Milestones (one per line)",
                               "\n".join(result.get("milestones", [])), height=90)
        rk_text = st.text_area("Risks (one per line)",
                               "\n".join(result.get("risks", [])), height=90)

        st.divider()
        t1, t2, t3 = st.columns([1.2, 1.6, 1.2])
        target = t1.radio("Apply to", ["New project", "Existing project"])
        projects = list_projects()
        existing = t2.selectbox("Existing project",
                                projects, format_func=lambda p: f"{p.code} — {p.name}") \
            if projects else None
        new_code = t2.text_input("New project code", "PRJ-KB-001") \
            if target == "New project" else ""
        overwrite = t3.checkbox("Overwrite non-empty fields", value=False,
                                help="Off = only fill fields that are currently empty.")

        if st.form_submit_button("Apply to project", type="primary"):
            edited.update({"name": name, "customer": customer,
                           "project_type": ptype, "region": region})
            p = _apply(edited,
                       [l for l in ms_text.splitlines() if l.strip()],
                       [l for l in rk_text.splitlines() if l.strip()],
                       "new" if target == "New project" else "existing",
                       existing.id if (existing and target == "Existing project") else None,
                       new_code, overwrite)
            if p is not None:
                st.session_state.pop("kb_result", None)
                st.success(f"Applied to **{p.code} — {p.name}**. Open Project "
                           "Initiation to review the charter, or the Delivery "
                           "Board to see it on the rail.")
