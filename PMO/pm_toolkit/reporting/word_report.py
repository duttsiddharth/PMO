"""Word (.docx) weekly status report generator (python-docx)."""
import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _heading(doc, text, size=14, color=(37, 99, 235)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def build_status_docx(report: dict) -> bytes:
    """report keys: project, customer, week, exec_summary, completed, upcoming,
    risks, issues, budget_status, schedule_status, decisions (lists/strings)."""
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"Weekly Status Report — {report.get('project','')}")
    r.bold = True
    r.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Customer: {report.get('customer','')}   |   "
                 f"Week ending: {report.get('week', date.today())}   |   "
                 f"Prepared by SD Advisory").italic = True

    _heading(doc, "Executive Summary")
    doc.add_paragraph(report.get("exec_summary", ""))

    def bullet_section(title, items):
        _heading(doc, title)
        if not items:
            doc.add_paragraph("None.", style="List Bullet")
            return
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    bullet_section("Completed This Period", report.get("completed", []))
    bullet_section("Planned Next Period", report.get("upcoming", []))
    bullet_section("Key Risks", report.get("risks", []))
    bullet_section("Open Issues", report.get("issues", []))

    _heading(doc, "Budget Status")
    doc.add_paragraph(report.get("budget_status", ""))
    _heading(doc, "Schedule Status")
    doc.add_paragraph(report.get("schedule_status", ""))
    bullet_section("Decisions Required", report.get("decisions", []))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
